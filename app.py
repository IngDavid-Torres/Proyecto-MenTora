import os
import io
import re
import time
import json
import random
from datetime import datetime


from flask import Flask, render_template, request, redirect, url_for, flash, session, jsonify
from flask_wtf import CSRFProtect
from flask_socketio import SocketIO, emit
from werkzeug.utils import secure_filename
from werkzeug.security import generate_password_hash, check_password_hash
from markupsafe import Markup


from config import SQLALCHEMY_DATABASE_URI, SECRET_KEY
from models import db, User, Quiz, Question, UserAnswer, Achievement, Badge, Notification, AccessLog, Teacher



app = Flask(__name__)
app.config['SQLALCHEMY_DATABASE_URI'] = SQLALCHEMY_DATABASE_URI
app.config['SECRET_KEY'] = SECRET_KEY
app.config['SITE_NAME'] = 'MenTora'
db.init_app(app)
socketio = SocketIO(app)

# Protección CSRF
csrf = CSRFProtect(app)

# Activar recarga automática
app.config['DEBUG'] = True
app.config['TEMPLATES_AUTO_RELOAD'] = True
app.jinja_env.auto_reload = True


# Función personalizada para manejar JSON correctamente
@app.template_filter('safe_json')
def safe_json_filter(data):
    """Convierte datos a JSON seguro para HTML"""
    if data is None:
        return '[]'
    try:
        # Usar separators para JSON compacto y ensure_ascii para evitar problemas
        json_str = json.dumps(data, ensure_ascii=True, separators=(',', ':'))
        return Markup(json_str)
    except Exception as e:
        print(f"Error en safe_json_filter: {e}")
        # Fallback: devolver array vacío
        return '[]'

@app.route('/chat')
def chat():
    return render_template('chat.html')



chat_history = []


@socketio.on('send_message')
def handle_send_message(data):
    username = data.get('username', 'Invitado')
    message = data.get('message', '')
    # Reemplazar atajos de emoji por emojis reales
    emoji_map = {
        ':)': '😊', ':(': '😢', ':D': '😃', '<3': '❤️', ':o': '😮', ':p': '😛', ':fire:': '🔥', ':star:': '⭐', ':ok:': '👌', ':cool:': '😎', ':rocket:': '🚀', ':100:': '💯', ':party:': '🥳', ':clap:': '👏', ':sad:': '😔', ':up:': '👍', ':down:': '👎', ':wink:': '😉', ':joy:': '😂', ':cry:': '😭', ':angry:': '😠', ':heart:': '❤️', ':check:': '✅', ':x:': '❌', ':star2:': '🌟', ':tada:': '🎉', ':wave:': '👋', ':smile:': '😄', ':sunglasses:': '😎', ':thinking:': '🤔', ':sleep:': '😴', ':zzz:': '💤', ':hug:': '🤗', ':pray:': '🙏', ':muscle:': '💪', ':eyes:': '👀', ':see_no_evil:': '🙈', ':poop:': '💩', ':cat:': '🐱', ':dog:': '🐶', ':robot:': '🤖', ':star-struck:': '🤩', ':mindblown:': '🤯', ':nerd:': '🤓', ':money:': '🤑', ':sweat:': '😅', ':kiss:': '😘', ':hugging:': '🤗', ':confetti:': '🎊', ':medal:': '🏅', ':trophy:': '🏆', ':crown:': '👑', ':medal2:': '🎖️', ':medal3:': '🥇', ':medal4:': '🥈', ':medal5:': '🥉', ':star3:': '⭐️', ':star4:': '🌠', ':star5:': '✨', ':star6:': '🌟', ':star7:': '💫', ':star8:': '🌟', ':star9:': '🌟', ':star10:': '🌟', ':star11:': '🌟', ':star12:': '🌟', ':star13:': '🌟', ':star14:': '🌟', ':star15:': '🌟', ':star16:': '🌟', ':star17:': '🌟', ':star18:': '🌟', ':star19:': '🌟', ':star20:': '🌟', ':star21:': '🌟', ':star22:': '🌟', ':star23:': '🌟', ':star24:': '🌟', ':star25:': '🌟', ':star26:': '🌟', ':star27:': '🌟', ':star28:': '🌟', ':star29:': '🌟', ':star30:': '🌟', ':star31:': '🌟', ':star32:': '🌟', ':star33:': '🌟', ':star34:': '🌟', ':star35:': '🌟', ':star36:': '🌟', ':star37:': '🌟', ':star38:': '🌟', ':star39:': '🌟', ':star40:': '🌟', ':star41:': '🌟', ':star42:': '🌟', ':star43:': '🌟', ':star44:': '🌟', ':star45:': '🌟', ':star46:': '🌟', ':star47:': '🌟', ':star48:': '🌟', ':star49:': '🌟', ':star50:': '🌟'
    }
    def replace_emojis(text):
        for k, v in emoji_map.items():
            text = text.replace(k, v)
        return text
    message = replace_emojis(message)
    chat_history.append({'username': username, 'message': message})
    if len(chat_history) > 50:
        chat_history.pop(0)
    emit('receive_message', {'username': username, 'message': message}, broadcast=True)

# Evento para enviar historial al conectarse
@socketio.on('request_history')
def handle_request_history():
    emit('chat_history', chat_history)


from werkzeug.security import generate_password_hash

with app.app_context():
    db.create_all()
    print("✅ Tablas creadas correctamente.")

    admin_password = os.environ.get('ADMIN_PASSWORD', 'admin12345!')
    print(f"[DEBUG] Contraseña admin usada: '{admin_password}'")
    # Eliminar usuario admin si existe
    admin = User.query.filter_by(username='admin').first()
    if admin:
        db.session.delete(admin)
        db.session.commit()
        print("⚠️ Usuario admin eliminado para recreación.")
    #
    admin_user = User(
        username='admin',
        email=None,
        password=generate_password_hash(admin_password),
        area='general',
        is_admin=True,
        points=0,
        level=1
    )
    db.session.add(admin_user)
    db.session.commit()
    print(f"👑 Usuario administrador creado (contraseña: {admin_password})")


@app.route('/')
def index():
    return render_template('index.html')

@app.route('/register', methods=['GET', 'POST'])
def register():
    if request.method == 'POST':
        username = request.form['username'].strip()
        email = request.form['email'].strip()
        password = request.form['password']
        area = request.form['area']
        user_type = request.form.get('user_type', 'student')

        if not username or not email or not password or not area or not user_type:
            msg = 'Todos los campos son obligatorios.'
            if request.accept_mimetypes['application/json']:
                return jsonify(success=False, message=msg), 400
            flash(msg, 'error')
            return redirect(url_for('register'))

        existing_user = User.query.filter_by(username=username).first()
        existing_email = User.query.filter_by(email=email).first()
        if existing_user:
            msg = 'El nombre de usuario ya está en uso.'
            if request.accept_mimetypes['application/json']:
                return jsonify(success=False, message=msg), 400
            flash(msg, 'error')
            return redirect(url_for('register'))
        if existing_email:
            msg = 'El correo electrónico ya está registrado.'
            if request.accept_mimetypes['application/json']:
                return jsonify(success=False, message=msg), 400
            flash(msg, 'error')
            return redirect(url_for('register'))

        try:
            hashed_password = generate_password_hash(password)
            new_user = User(
                username=username,
                email=email,
                password=hashed_password,
                area=area,
                points=0,
                level=1,
                is_admin=False
            )
            db.session.add(new_user)
            db.session.flush()  # Para obtener el id
            if user_type == 'teacher':
                teacher = Teacher(user_id=new_user.id, area=area)
                db.session.add(teacher)
            db.session.commit()
            msg = 'Registro exitoso. Ahora puedes iniciar sesión.'
            if request.accept_mimetypes['application/json']:
                return jsonify(success=True, message=msg)
            flash(msg, 'success')
            return redirect(url_for('login'))
        except Exception as e:
            db.session.rollback()
            app.logger.exception('Error al registrar usuario:')
            # Devolver un mensaje de error correcto (no reutilizar el mensaje de éxito)
            err_msg = 'Error al registrar usuario. Intenta de nuevo más tarde.'
            if request.accept_mimetypes['application/json']:
                return jsonify(success=False, message=err_msg), 500
            flash(err_msg, 'error')
            return redirect(url_for('register'))

    return render_template('register.html')

@app.route('/login', methods=['GET', 'POST'])
def login():

    if request.method == 'POST':
        username = request.form.get('username', '').strip()
        password = request.form.get('password', '')
        app.logger.info(f"[LOGIN DEBUG] username recibido: '{username}' | password recibido: '{password}'")
        if not username or not password:
            app.logger.error(f"[LOGIN ERROR] Campos faltantes: username='{username}', password='{password}'")
            return jsonify(success=False, error="Campos requeridos faltantes"), 400
        user = User.query.filter_by(username=username).first()
        ip = request.remote_addr or 'unknown'
        success = user is not None and check_password_hash(user.password, password)
        app.logger.info(f"[LOGIN DEBUG] Usuario encontrado: {user is not None} | Password match: {success}")
        log = AccessLog(username=username, ip=ip, success=success)
        db.session.add(log)
        db.session.commit()
        if success:
            session['user_id'] = user.id
            session['username'] = user.username
            session['is_admin'] = user.is_admin
            session['avatar_url'] = user.avatar_url if user.avatar_url else ''
            session['theme'] = user.theme if user.theme else 'default'
            # Redirección según tipo de usuario
            if user.is_admin:
                return jsonify(success=True, redirect=url_for('admin_panel'))
            elif hasattr(user, 'teacher_profile') and user.teacher_profile is not None:
                return jsonify(success=True, redirect=url_for('teacher_dashboard'))
            else:
                return jsonify(success=True, redirect=url_for('dashboard'))
        else:
            app.logger.error(f"[LOGIN ERROR] Usuario o contraseña incorrectos para username='{username}'")
            return jsonify(success=False, error="Usuario o contraseña incorrectos"), 401

    
    return render_template('login.html')


@app.route('/dashboard')
def dashboard():
    if 'user_id' not in session:
        return redirect(url_for('login'))

    user = User.query.get(session['user_id'])

   
    from datetime import date
    import random
    daily_seed = int(date.today().strftime('%Y%m%d'))
    quizzes_in_area = Quiz.query.filter_by(area=user.area).all()
    daily_quiz = None
    if quizzes_in_area:
        random.seed(daily_seed)
        daily_quiz = random.choice(quizzes_in_area)
    from sqlalchemy.orm import joinedload
    achievements = Achievement.query.options(joinedload(Achievement.badge)).filter_by(user_id=user.id).all()
    
    admin_user = User.query.filter_by(is_admin=True).first()
   
    from models import Teacher
    teacher = None
    if hasattr(user, 'teacher_profile') and user.teacher_profile is not None:
        teacher = user.teacher_profile
    else:
        teacher = Teacher.query.filter_by(user_id=user.id).first()

    from sqlalchemy import or_
    quizzes_all = Quiz.query.filter(
            or_(Quiz.area == user.area, Quiz.teacher_id != None, Quiz.teacher_id == None)
    ).all()
    quizzes = []
    for q in quizzes_all:
        tiene_preguntas = hasattr(q, 'questions') and q.questions and len(q.questions) > 0
        es_del_profesor = teacher is not None and hasattr(q, 'teacher_id') and q.teacher_id == teacher.id
        es_del_admin = (
            (hasattr(q, 'teacher') and getattr(q.teacher, 'is_admin', False)) or
            (q.teacher_id is None)
        )
        if tiene_preguntas and (es_del_profesor or es_del_admin):
            quizzes.append(q)
    from models import Game
    games_all = Game.query.filter(
        or_(Game.area == user.area, Game.teacher_id == None)
    ).all() if hasattr(Game, 'area') else []
    games = []
    for g in games_all:
        es_del_profesor = teacher is not None and hasattr(g, 'teacher_id') and g.teacher_id == teacher.id
        es_del_admin = (
            (hasattr(g, 'teacher') and getattr(g.teacher, 'is_admin', False)) or
            (g.teacher_id is None)
        )
        if es_del_profesor or es_del_admin:
            games.append(g)
    
    last_answers = (
        db.session.query(UserAnswer, Question, Quiz)
        .join(Question, UserAnswer.question_id == Question.id)
        .join(Quiz, Question.quiz_id == Quiz.id)
        .filter(UserAnswer.user_id == user.id)
        .order_by(UserAnswer.answered_at.desc())
        .limit(5)
        .all()
    )
    last_achievements = Achievement.query.filter_by(user_id=user.id).order_by(Achievement.earned_at.desc()).limit(5).all()

    
    points_this_level = user.points - ((user.level - 1) * 100)
    points_needed = 100
    progress_percent = int((points_this_level / points_needed) * 100) if user.level > 0 else 0
    if progress_percent > 100:
        progress_percent = 100
        # Calcular progreso considerando quizzes de profesor con 15 puntos y admin con 10
        points_this_level = 0
        quizzes_answered_ids = db.session.query(Question.quiz_id).join(UserAnswer).filter(UserAnswer.user_id == user.id).distinct().all()
        quizzes_answered = set(qid for (qid,) in quizzes_answered_ids)
        for q in quizzes:
            if q.id in quizzes_answered:
                if q.teacher_id:
                    points_this_level += 15
                else:
                    points_this_level += 10
        points_this_level += user.points - ((user.level - 1) * 100)

   
    if progress_percent == 100:
        motivational_message = "¡Felicidades! Has alcanzado un nuevo nivel. ¡Sigue así!"
    elif progress_percent >= 75:
        motivational_message = "¡Estás muy cerca de subir de nivel!"
    elif progress_percent >= 50:
        motivational_message = "¡Vas por la mitad, no te detengas!"
    elif progress_percent >= 25:
        motivational_message = "¡Buen progreso! Sigue participando para subir de nivel."
    else:
        motivational_message = "¡Cada reto cuenta! Responde más para avanzar."

    # Ranking: top 10 usuarios por puntos
    leaderboard = User.query.filter_by(is_admin=False).order_by(User.points.desc(), User.level.desc()).limit(10).all()


    # Progreso de retos completados
    total_quizzes = Quiz.query.filter_by(area=user.area).count()
    quizzes_answered_ids = db.session.query(Question.quiz_id).join(UserAnswer).filter(UserAnswer.user_id == user.id).distinct().all()
    quizzes_answered = len(set(qid for (qid,) in quizzes_answered_ids))
    if total_quizzes > 0:
        quiz_progress_percent = int((quizzes_answered / total_quizzes) * 100)
    else:
        quiz_progress_percent = 0

    # Notificaciones: mostrar solo las globales y las dirigidas a este usuario
    notifications = Notification.query.filter(
        (Notification.to_user == None) | (Notification.to_user == user.id)
    ).order_by(Notification.date_sent.desc()).limit(5).all()

    return render_template('dashboard.html',
                           username=user.username,
                           area=user.area,
                           points=user.points,
                           level=user.level,
                           achievements=achievements,
                           quizzes=quizzes,
                           games=games,
                           progress_percent=progress_percent,
                           motivational_message=motivational_message,
                           leaderboard=leaderboard,
                           total_quizzes=total_quizzes,
                           quizzes_answered=quizzes_answered,
                           quiz_progress_percent=quiz_progress_percent,
                           last_answers=last_answers,
                           last_achievements=last_achievements,
                           notifications=notifications,
                           avatar_url=user.avatar_url,
                           theme=session.get('theme', user.theme if user.theme else 'default'),
                           daily_quiz=daily_quiz,
                           teacher=teacher)

@app.route('/biblioteca')
def biblioteca():
    """Biblioteca digital con recursos de aprendizaje."""
    if 'user_id' not in session:
        return redirect(url_for('login'))
        
    try:
        user = User.query.get(session['user_id'])
        if not user:
            flash('Usuario no encontrado.', 'error')
            return redirect(url_for('login'))
        
        return render_template('biblioteca.html',
                             username=user.username,
                             avatar_url=user.avatar_url,
                             theme=session.get('theme', user.theme if user.theme else 'default'))
    except Exception as e:
        app.logger.error(f"Error en biblioteca: {str(e)}")
        flash('Error al cargar la biblioteca.', 'error')
        return redirect(url_for('dashboard'))
    
    
# Reto activo
@app.route('/quiz', methods=['GET', 'POST'])
def quiz():
    if 'user_id' not in session:
        return redirect(url_for('login'))

    user = User.query.get(session['user_id'])
    # Si es GET y no se está intentando un quiz específico, mostrar lista de quizzes disponibles
    from sqlalchemy import or_
    quizzes_all = Quiz.query.filter(
        or_(Quiz.area == user.area, Quiz.teacher_id == None)
    ).all()
    quizzes = []
    for q in quizzes_all:
        tiene_preguntas = hasattr(q, 'questions') and q.questions and len(q.questions) > 0
        if tiene_preguntas:
            quizzes.append(q)
    # Si no es POST, mostrar lista de quizzes
    if request.method == 'GET' and not request.args.get('quiz_id'):
        return render_template('quiz.html', all_quizzes=quizzes)

    # Si es POST o se está intentando un quiz aleatorio (legacy)
    question = Question.query.join(Quiz).filter(Quiz.area == user.area).order_by(db.func.random()).first()
    if not question:
        flash('No hay preguntas disponibles en tu área aún.')
        return redirect(url_for('dashboard'))

    if request.method == 'POST':
        selected = request.form.get('selected_answer')
        is_correct = selected == question.correct_answer
        answer = UserAnswer(
            user_id=user.id,
            question_id=question.id,
            selected_answer=selected,
            is_correct=is_correct,
            answered_at=datetime.utcnow()
        )
        db.session.add(answer)
        if is_correct:
            user.points += 10
            if user.points >= user.level * 100:
                user.level += 1

            # Otorgar logro por primer quiz correcto
            from models import Achievement, Badge
            primer_quiz_badge = Badge.query.filter_by(name='Primer Quiz').first()
            if primer_quiz_badge:
                ya_tiene = Achievement.query.filter_by(user_id=user.id, badge_id=primer_quiz_badge.id).first()
                if not ya_tiene:
                    nuevo_logro = Achievement(user_id=user.id, badge_id=primer_quiz_badge.id)
                    db.session.add(nuevo_logro)
        db.session.commit()

        flash('Respuesta registrada. ' + ('¡Correcto!' if is_correct else 'Incorrecto.'))
        return redirect(url_for('dashboard'))

    return render_template('quiz.html', question=question)


@app.route('/noticias')
def noticias():
    """Noticiero de programación con actualizaciones tech."""
    if 'user_id' not in session:
        return redirect(url_for('login'))
        
    try:
        user = User.query.get(session['user_id'])
        if not user:
            flash('Usuario no encontrado.', 'error')
            return redirect(url_for('login'))
        
        return render_template('noticias.html',
                             username=user.username,
                             avatar_url=user.avatar_url,
                             theme=session.get('theme', user.theme if user.theme else 'default'))
    except Exception as e:
        app.logger.error(f"Error en noticias: {str(e)}")
        flash('Error al cargar las noticias.', 'error')
        return redirect(url_for('dashboard'))
    



@app.route('/admin', methods=['GET'])
def admin_panel():
    if 'user_id' not in session or not session.get('is_admin'):
        flash('Acceso restringido. Solo administradores.')
        return redirect(url_for('login'))

    total_users = User.query.count()
    total_quizzes = Quiz.query.count()
    total_questions = Question.query.count()
    users = User.query.filter(User.username != 'admin').all()
    user_to_edit = session.pop('user_to_edit', None)
    from models import Game
    games = Game.query.order_by(Game.created_at.desc()).all()
    game_to_edit = session.get('game_to_edit', None)
    quizzes = Quiz.query.order_by(Quiz.created_at.desc()).all()
    achievements = Badge.query.order_by(Badge.created_at.desc()).all()
    achievement_to_edit = session.pop('achievement_to_edit', None)
    # Notificaciones: últimas 20
    notifications = Notification.query.order_by(Notification.date_sent.desc()).limit(20).all()
    # Logs de acceso: últimos 20
    access_logs = AccessLog.query.order_by(AccessLog.timestamp.desc()).limit(20).all()
    return render_template('admin.html',
                           username=session['username'],
                           total_users=total_users,
                           total_quizzes=total_quizzes,
                           total_questions=total_questions,
                           users=users,
                           user_to_edit=user_to_edit,
                           games=games,
                           game_to_edit=game_to_edit,
                           quizzes=quizzes,
                           achievements=achievements,
                           achievement_to_edit=achievement_to_edit,
                           notifications=notifications,
                           access_logs=access_logs)
# Cambiar contraseña de admin
@app.route('/admin/change_password', methods=['POST'])
def change_admin_password():
    if 'user_id' not in session or not session.get('is_admin'):
        flash('Acceso restringido.')
        return redirect(url_for('login'))
    current_password = request.form.get('current_password')
    new_password = request.form.get('new_password')
    confirm_password = request.form.get('confirm_password')
    admin = User.query.get(session['user_id'])
    if not check_password_hash(admin.password, current_password):
        flash('La contraseña actual es incorrecta.')
        return redirect(url_for('admin_panel'))
    if new_password != confirm_password:
        flash('La nueva contraseña y la confirmación no coinciden.')
        return redirect(url_for('admin_panel'))
    if len(new_password) < 6:
        flash('La nueva contraseña debe tener al menos 6 caracteres.')
        return redirect(url_for('admin_panel'))
    admin.password = generate_password_hash(new_password)
    db.session.commit()
    flash('Contraseña actualizada correctamente.')
    return redirect(url_for('admin_panel'))


@app.route('/admin/send_notification', methods=['POST'])
def send_notification():
    if 'user_id' not in session or not session.get('is_admin'):
        flash('Acceso restringido.')
        return redirect(url_for('login'))
    message = request.form.get('notification_message', '').strip()
    to_user = request.form.get('notification_user')
    if not message:
        flash('El mensaje no puede estar vacío.')
        return redirect(url_for('admin_panel'))
    if to_user == 'all':
        notif = Notification(message=message, to_user=None)
        db.session.add(notif)
    else:
        try:
            user_id = int(to_user)
            notif = Notification(message=message, to_user=user_id)
            db.session.add(notif)
        except Exception:
            flash('Usuario destino inválido.')
            return redirect(url_for('admin_panel'))
    db.session.commit()
    flash('Notificación enviada.')
    return redirect(url_for('admin_panel'))


@app.route('/admin/create_or_update_achievement', methods=['POST'])
def create_or_update_achievement():
    if 'user_id' not in session or not session.get('is_admin'):
        flash('Acceso restringido.')
        return redirect(url_for('login'))
    achievement_id = request.form.get('achievement_id')
    name = request.form['achievement_name'].strip()
    description = request.form['achievement_description'].strip()
    image_file = request.files.get('achievement_image')
    image_url = None
    if image_file and image_file.filename:
        filename = secure_filename(image_file.filename)
        img_folder = os.path.join('static', 'img', 'badges')
        os.makedirs(img_folder, exist_ok=True)
        img_path = os.path.join(img_folder, filename)
        image_file.save(img_path)
        image_url = f'/static/img/badges/{filename}'
    if achievement_id:
        badge = Badge.query.get(achievement_id)
        if badge:
            badge.name = name
            badge.description = description
            if image_url:
                badge.image_url = image_url
            db.session.commit()
            flash('Logro/insignia actualizado.')
    else:
        badge = Badge(name=name, description=description, image_url=image_url)
        db.session.add(badge)
        db.session.commit()
        flash('Logro/insignia creado.')
    return redirect(url_for('admin_panel'))


@app.route('/admin/edit_achievement/<int:badge_id>', methods=['GET'])
def edit_achievement(badge_id):
    if 'user_id' not in session or not session.get('is_admin'):
        flash('Acceso restringido.')
        return redirect(url_for('login'))
    badge = Badge.query.get_or_404(badge_id)
    session['achievement_to_edit'] = {
        'id': badge.id,
        'name': badge.name,
        'description': badge.description,
        'image_url': badge.image_url
    }
    return redirect(url_for('admin_panel'))


@app.route('/admin/clear_achievement_edit')
def clear_achievement_edit():
    session.pop('achievement_to_edit', None)
    return redirect(url_for('admin_panel'))


@app.route('/admin/delete_achievement/<int:badge_id>', methods=['POST'])
def delete_achievement(badge_id):
    if 'user_id' not in session or not session.get('is_admin'):
        flash('Acceso restringido.')
        return redirect(url_for('login'))
    badge = Badge.query.get_or_404(badge_id)
    db.session.delete(badge)
    db.session.commit()
    flash('Logro/insignia eliminado.')
    return redirect(url_for('admin_panel'))
# Ruta para crear juego desde admin
@app.route('/admin/create_game', methods=['POST'])

def create_game():
    if 'user_id' not in session or not session.get('is_admin'):
        flash('Acceso restringido.')
    import os
    admin_password = os.environ.get('ADMIN_PASSWORD', 'admin12345!')
    admin = User.query.filter_by(username='admin').first()
    if admin:
        admin.password = generate_password_hash(admin_password)
        admin.area = 'general'
        admin.is_admin = True
        db.session.commit()
        print("🔐 Usuario admin actualizado (contraseña por variable de entorno).")
    else:
        admin_user = User(
            username='admin',
            email=None,
            password=generate_password_hash(admin_password),
            area='general',
            is_admin=True,
            points=0,
            level=1
        )
        db.session.add(admin_user)
        db.session.commit()
        print("👑 Usuario administrador creado (contraseña por variable de entorno).")
    
    # Limpiar datos de edición
    session.pop('game_to_edit', None)
    return redirect(url_for('admin_panel'))

# Eliminar juego
@app.route('/admin/delete_game/<int:game_id>', methods=['POST'])
def delete_game(game_id):
    if 'user_id' not in session or not session.get('is_admin'):
        flash('Acceso restringido.')
        return redirect(url_for('login'))
    
    from models import Game
    game = Game.query.get_or_404(game_id)
    
    try:
        db.session.delete(game)
        db.session.commit()
        flash(f'Juego "{game.name}" eliminado correctamente.')
    except Exception as e:
        db.session.rollback()
        flash(f'Error al eliminar el juego: {e}')
    
    return redirect(url_for('admin_panel'))

# Editar juego (cargar en formulario)
@app.route('/admin/edit_game/<int:game_id>', methods=['GET'])
def edit_game(game_id):
    if 'user_id' not in session or not session.get('is_admin'):
        flash('Acceso restringido.')
        return redirect(url_for('login'))
    
    from models import Game
    game = Game.query.get_or_404(game_id)
    session['game_to_edit'] = {
        'id': game.id,
        'name': game.name,
        'description': game.description,
        'rules': game.rules
    }
    return redirect(url_for('admin_panel'))

# Cancelar edición de juego
@app.route('/admin/clear_game_edit')
def clear_game_edit():
    session.pop('game_to_edit', None)
    return redirect(url_for('admin_panel'))

# Ruta para crear quiz desde admin
@app.route('/admin/create_quiz', methods=['POST'])

def create_quiz():
    if 'user_id' not in session or not session.get('is_admin'):
        flash('Acceso restringido.')
        return redirect(url_for('login'))

    title = request.form.get('quiz_name', '').strip()
    area = request.form.get('quiz_area', '').strip()
    description = request.form.get('quiz_description', '').strip()

    if not title or not area or not description:
        flash('Todos los campos son obligatorios.')
        return redirect(url_for('admin_panel'))

    # Si el usuario es admin, teacher_id debe ser None
    new_quiz = Quiz(
        title=title,
        area=area,
        description=description,
        created_by=session['user_id'],
        teacher_id=None,
        created_at=datetime.utcnow()
    )
    db.session.add(new_quiz)
    db.session.commit()
    flash(f'Quiz "{title}" creado exitosamente.')
    return redirect(url_for('admin_panel'))

@app.route('/logout', methods=['GET', 'POST'])
def logout():
    session.clear()
    # Mark logout message as an 'error' style (red) to match the site's visual language
    flash('Sesión cerrada correctamente.', 'error')
    response = redirect(url_for('login'))
    response.headers['Cache-Control'] = 'no-store, no-cache, must-revalidate, max-age=0, private'
    response.headers['Pragma'] = 'no-cache'
    response.headers['Expires'] = '0'
    # Invalida cookies de sesión
    response.set_cookie('session', '', expires=0)
    return response

@app.route('/edit_profile', methods=['GET'])
def edit_profile():
    if 'user_id' not in session:
        return redirect(url_for('login'))

    user = User.query.get(session['user_id'])
    # Opciones de temas y avatares
    themes = ['default', 'dark', 'light', 'blue', 'green', 'pink']
    avatar_seeds = ['default', 'cat', 'dog', 'robot', 'coder', 'star', 'rocket', 'owl', 'fox', 'lion', 'panda', 'alien', 'ninja', 'wizard', 'unicorn', 'dragon']
    return render_template('edit_profile.html', user=user, themes=themes, avatar_seeds=avatar_seeds)


@app.route('/update_profile', methods=['POST'])
def update_profile():
    if 'user_id' not in session:
        return redirect(url_for('login'))

    user = User.query.get(session['user_id'])

    new_email = request.form['email'].strip()
    new_area = request.form['area']
    new_avatar_seed = request.form.get('avatar_seed')
    new_theme = request.form.get('theme')

    if new_email:
        user.email = new_email
    if new_area:
        user.area = new_area
    if new_avatar_seed:
        user.avatar_url = f"https://api.dicebear.com/7.x/identicon/svg?seed={new_avatar_seed}"
    if new_theme:
        user.theme = new_theme

    db.session.commit()
    # Actualizar sesión
    session['avatar_url'] = user.avatar_url if user.avatar_url else ''
    session['theme'] = user.theme if user.theme else 'default'
    if request.headers.get('Accept') == 'application/json':
        return jsonify(success=True, theme=session['theme'])
    flash('Perfil actualizado correctamente.')
    return redirect(url_for('dashboard'))



@app.route('/games')
def games():
    if 'user_id' not in session:
        return redirect(url_for('login'))
    from models import User, Quiz, Game
    from sqlalchemy import or_
    user = User.query.get(session['user_id'])
    # Quizzes igual que en dashboard
    quizzes_all = Quiz.query.filter(
        or_(Quiz.area == user.area, Quiz.teacher_id == None)
    ).all()
    quizzes = []
    for q in quizzes_all:
        tiene_preguntas = hasattr(q, 'questions') and q.questions and len(q.questions) > 0
        if tiene_preguntas:
            quizzes.append(q)
    # Juegos igual que en dashboard
    games_all = Game.query.filter(
        or_(Game.area == user.area, Game.teacher_id == None)
    ).all() if hasattr(Game, 'area') else Game.query.order_by(Game.name).all()
    games = []
    for g in games_all:
        games.append(g)
    return render_template('games.html', quizzes=quizzes, games=games)

@app.route('/juegos-interactivos')
def juegos_interactivos():
    """Ruta para los juegos interactivos (programación y VR)"""
    if 'user_id' not in session:
        return redirect(url_for('login'))
    
    user = User.query.get(session['user_id'])
    return render_template('juegos_interactivos.html', 
                         username=user.username,
                         user_id=user.id,
                         avatar_url=user.avatar_url,
                         theme=session.get('theme', user.theme if user.theme else 'default'))


@app.route('/games/<int:game_id>/run', methods=['POST'])
def run_game_code(game_id):
    from models import Game, User
    import io, contextlib, time
    game = Game.query.get_or_404(game_id)
    code = request.form.get('code', '')
    output = ''
    feedback = ''
    max_attempts = 3
    cooldown_minutes = 30
    user_id = session.get('user_id')
    if not user_id:
        return redirect(url_for('login'))
    user = User.query.get(user_id)
    # Guardar intentos y cooldown en sesión por usuario y juego
    session_key = f'game_{game_id}_attempts'
    session_cooldown_key = f'game_{game_id}_cooldown'
    attempts = session.get(session_key, 0)
    cooldown_until = session.get(session_cooldown_key)
    import datetime
    now = int(time.time())
    # Si está en cooldown
    if cooldown_until and now < cooldown_until:
        mins_left = int((cooldown_until - now) / 60) + 1
        feedback = f'Sin vidas. Intenta de nuevo en {mins_left} minutos.'
        return render_template('game_detail.html', game=game, output=output, feedback=feedback, attempts=max_attempts, max_attempts=max_attempts, cooldown=True)
    # Si no está en cooldown, resetear si ya pasó
    if cooldown_until and now >= cooldown_until:
        attempts = 0
        session[session_key] = 0
        session.pop(session_cooldown_key, None)
    # Solo permitir si tiene intentos
    if attempts >= max_attempts:
        session[session_cooldown_key] = now + cooldown_minutes * 60
        feedback = f'Sin vidas. Intenta de nuevo en {cooldown_minutes} minutos.'
        return render_template('game_detail.html', game=game, output=output, feedback=feedback, attempts=max_attempts, max_attempts=max_attempts, cooldown=True)
    # Evaluar código
    if game.name.lower() == 'hola mundo':
        safe_builtins = {'print': print}
        buf = io.StringIO()
        try:
            with contextlib.redirect_stdout(buf):
                exec(code, {'__builtins__': safe_builtins})
            output = buf.getvalue()
        except Exception as e:
            output = f'Error: {str(e)}'
        # Validar si acertó (print exacto)
        if output.strip() == 'Hola mundo':
            feedback = '¡Correcto! Ganaste +10 puntos.'
            
            progress_percent = session.get('progress_percent', None)
            if progress_percent is None:
                
                points_this_level = user.points - ((user.level - 1) * 100)
                points_needed = 100
                progress_percent = int((points_this_level / points_needed) * 100) if user.level > 0 else 0
            progress_percent += 10
            user.points += 10
            
            if progress_percent >= 100 or user.points >= user.level * 100:
                progress_percent = 0
                user.level += 1
                feedback = '¡Excelente, así se hace! ¡Subes de nivel!'
            session['progress_percent'] = progress_percent
            db.session.commit()
            # Resetear intentos tras éxito
            session[session_key] = 0
            session.pop(session_cooldown_key, None)
            return render_template('game_detail.html', game=game, output=output, feedback=feedback, attempts=0, max_attempts=max_attempts, success=True, show_animated_alert=True)
        else:
            attempts += 1
            session[session_key] = attempts
            if attempts >= max_attempts:
                session[session_cooldown_key] = now + cooldown_minutes * 60
                feedback = f'Incorrecto. Sin vidas. Intenta de nuevo en {cooldown_minutes} minutos.'
                return render_template('game_detail.html', game=game, output=output, feedback=feedback, attempts=attempts, max_attempts=max_attempts, cooldown=True)
            else:
                feedback = f'Incorrecto. Intentos restantes: {max_attempts - attempts}'
    elif game.name.lower() == 'la suma de dos números':
        safe_builtins = {'print': print}
        buf = io.StringIO()
        try:
            with contextlib.redirect_stdout(buf):
                exec(code, {'__builtins__': safe_builtins})
            output = buf.getvalue()
        except Exception as e:
            output = f'Error: {str(e)}'
        import re
        nums = re.findall(r'(\d+)', code)
        if len(nums) >= 2:
            a, b = int(nums[0]), int(nums[1])
            suma_esperada = a + b
            try:
                salida = int(output.strip())
            except Exception:
                salida = None
            if salida == suma_esperada:
                feedback = '¡Correcto! Ganaste +10 puntos.'
                progress_percent = session.get('progress_percent', None)
                if progress_percent is None:
                    points_this_level = user.points - ((user.level - 1) * 100)
                    points_needed = 100
                    progress_percent = int((points_this_level / points_needed) * 100) if user.level > 0 else 0
                progress_percent += 10
                user.points += 10
                if progress_percent >= 100 or user.points >= user.level * 100:
                    progress_percent = 0
                    user.level += 1
                    feedback = '¡Excelente, así se hace! ¡Subes de nivel!'
                session['progress_percent'] = progress_percent
                db.session.commit()
                session[session_key] = 0
                session.pop(session_cooldown_key, None)
                return render_template('game_detail.html', game=game, output=output, feedback=feedback, attempts=0, max_attempts=max_attempts, success=True, show_animated_alert=True)
            else:
                attempts += 1
                session[session_key] = attempts
                if attempts >= max_attempts:
                    session[session_cooldown_key] = now + cooldown_minutes * 60
                    feedback = f'Incorrecto. Sin vidas. Intenta de nuevo en {cooldown_minutes} minutos.'
                    return render_template('game_detail.html', game=game, output=output, feedback=feedback, attempts=attempts, max_attempts=max_attempts, cooldown=True)
                else:
                    feedback = f'Incorrecto. Intentos restantes: {max_attempts - attempts}'
        else:
            feedback = 'Por favor, define dos números en tu código.'
    return render_template('game_detail.html', game=game, output=output, feedback=feedback, attempts=attempts, max_attempts=max_attempts)
@app.route('/games/<int:game_id>')
def game_detail(game_id):
    from models import Game
    game = Game.query.get_or_404(game_id)
    # Para mostrar vidas aunque sea GET
    max_attempts = 3
    session_key = f'game_{game_id}_attempts'
    session_cooldown_key = f'game_{game_id}_cooldown'
    attempts = session.get(session_key, 0)
    import time
    now = int(time.time())
    cooldown_until = session.get(session_cooldown_key)
    cooldown = False
    if cooldown_until and now < cooldown_until:
        cooldown = True
    return render_template('game_detail.html', game=game, attempts=attempts, max_attempts=max_attempts, cooldown=cooldown)

@app.route('/chatbot', methods=['POST'])
def chatbot():
    data = request.get_json()
    user_msg = data.get('message', '').strip()
    # Lógica simple: eco o respuesta programada
    if not user_msg:
        return jsonify({'response': '¿Puedes escribir tu pregunta?'}), 200
   
    if 'hola' in user_msg.lower():
        return jsonify({'response': '¡Hola! ¿En qué puedo ayudarte hoy?'}), 200
    if 'nivel' in user_msg.lower():
        return jsonify({'response': 'Puedes subir de nivel completando retos y juegos.'}), 200
    if 'juego' in user_msg.lower():
        return jsonify({'response': 'Para jugar, ve a la sección de Juegos y elige uno.'}), 200
    # Eco por defecto
    return jsonify({'response': f'Recibí: {user_msg}'}), 200


@app.route('/quiz/<int:quiz_id>', methods=['GET', 'POST'])
def quiz_attempt(quiz_id):
    if 'user_id' not in session:
        return redirect(url_for('login'))
    user = User.query.get(session['user_id'])
    quiz = Quiz.query.get_or_404(quiz_id)
    # Obtener preguntas del quiz
    questions = Question.query.filter_by(quiz_id=quiz.id).all()
    if not questions:
        flash('Este quiz no tiene preguntas aún.')
        return redirect(url_for('dashboard'))
    max_attempts = 3
    cooldown_minutes = 30
    session_key = f'quiz_{quiz_id}_attempts'
    session_cooldown_key = f'quiz_{quiz_id}_cooldown'
    attempts = session.get(session_key, 0)
    cooldown_until = session.get(session_cooldown_key)
    import time
    now = int(time.time())
    if cooldown_until and now < cooldown_until:
        mins_left = int((cooldown_until - now) / 60) + 1
        return render_template('quiz.html', quiz=quiz, questions=questions, attempts=max_attempts, max_attempts=max_attempts, cooldown=True, mins_left=mins_left)
    if cooldown_until and now >= cooldown_until:
        attempts = 0
        session[session_key] = 0
        session.pop(session_cooldown_key, None)
    if attempts >= max_attempts:
        session[session_cooldown_key] = now + cooldown_minutes * 60
        mins_left = cooldown_minutes
        return render_template('quiz.html', quiz=quiz, questions=questions, attempts=max_attempts, max_attempts=max_attempts, cooldown=True, mins_left=mins_left)
    if request.method == 'POST':
        correctas = 0
        total = len(questions)
        user_answers = {}
        for q in questions:
            selected = request.form.get(f'q_{q.id}')
            user_answers[q.id] = selected
            is_correct = selected == q.correct_answer
            # Guardar respuesta
            answer = UserAnswer(
                user_id=user.id,
                question_id=q.id,
                selected_answer=selected,
                is_correct=is_correct,
                answered_at=datetime.utcnow()
            )
            db.session.add(answer)
            if is_correct:
                correctas += 1
        db.session.commit()
        results = [(q, user_answers.get(q.id)) for q in questions]
        if correctas == total:
            # Respuestas correctas, sumar puntos y resetear intentos
            # Si el quiz es de profesor, 15 puntos por pregunta, si no, 10
            puntos = 15 * total if quiz.teacher_id else 10 * total
            user.points += puntos
            db.session.commit()
            session[session_key] = 0
            session.pop(session_cooldown_key, None)
            show_15pts_modal = True
            return render_template('quiz.html', quiz=quiz, questions=questions, show_answers=True, results=results, attempts=0, max_attempts=max_attempts, show_15pts_modal=show_15pts_modal)
        else:
            # Respuestas incorrectas, aumentar intentos
            attempts += 1
            session[session_key] = attempts
            if attempts >= max_attempts:
                session[session_cooldown_key] = now + cooldown_minutes * 60
                mins_left = cooldown_minutes
                return render_template('quiz.html', quiz=quiz, questions=questions, attempts=max_attempts, max_attempts=max_attempts, cooldown=True, mins_left=mins_left)
            return render_template('quiz.html', quiz=quiz, questions=questions, attempts=attempts, max_attempts=max_attempts, error=True)
    return render_template('quiz.html', quiz=quiz, questions=questions, attempts=attempts, max_attempts=max_attempts)

@app.route('/admin/add_question', methods=['POST'])
def add_question():
    if 'user_id' not in session or not session.get('is_admin'):
        flash('Acceso restringido.')
        return redirect(url_for('login'))
    quiz_id = request.form['quiz_id']
    text = request.form['question_text'].strip()
    options = [
        request.form['option1'].strip(),
        request.form['option2'].strip(),
        request.form['option3'].strip(),
        request.form['option4'].strip()
    ]
    correct_index = int(request.form['correct_answer']) - 1
    if not text or any(not o for o in options) or correct_index not in range(4):
        flash('Todos los campos son obligatorios y la respuesta debe ser válida.')
        return redirect(url_for('admin_panel'))
    correct_answer = options[correct_index]
    new_question = Question(
        quiz_id=quiz_id,
        text=text,
        options=options,
        correct_answer=correct_answer
    )
    db.session.add(new_question)
    db.session.commit()
    flash('Pregunta agregada correctamente.')
    return redirect(url_for('admin_panel'))

@app.route('/admin/create_or_update_user', methods=['POST'])
def create_or_update_user():
    if 'user_id' not in session or not session.get('is_admin'):
        flash('Acceso restringido.')
        return redirect(url_for('login'))
    user_id = request.form.get('user_id')
    username = request.form['username'].strip()
    email = request.form['email'].strip()
    area = request.form['area'].strip()
    points = int(request.form.get('points', 0))
    level = int(request.form.get('level', 1))
    password = request.form.get('password')
    if user_id:
        user = User.query.get(user_id)
        if not user:
            flash('Usuario no encontrado.')
            return redirect(url_for('admin_panel'))
        user.username = username
        user.email = email
        user.area = area
        user.points = points
        user.level = level
        if password:
            user.password = generate_password_hash(password)
        db.session.commit()
        flash('Usuario actualizado correctamente.')
    else:
        if not password:
            flash('La contraseña es obligatoria para crear un usuario.')
            return redirect(url_for('admin_panel'))
        if User.query.filter_by(username=username).first():
            flash('El nombre de usuario ya existe.')
            return redirect(url_for('admin_panel'))
        if User.query.filter_by(email=email).first():
            flash('El email ya está registrado.')
            return redirect(url_for('admin_panel'))
        new_user = User(
            username=username,
            email=email,
            area=area,
            points=points,
            level=level,
            password=generate_password_hash(password),
            is_admin=False
        )
        db.session.add(new_user)
        db.session.commit()
        flash('Usuario creado correctamente.')
    session.pop('user_to_edit', None)
    return redirect(url_for('admin_panel'))

# Editar usuario (cargar en formulario)
@app.route('/admin/edit_user/<int:user_id>', methods=['GET'])
def edit_user(user_id):
    if 'user_id' not in session or not session.get('is_admin'):
        flash('Acceso restringido.')
        return redirect(url_for('login'))
    user = User.query.get_or_404(user_id)
    session['user_to_edit'] = {
        'id': user.id,
        'username': user.username,
        'email': user.email,
        'area': user.area,
        'points': user.points,
        'level': user.level
    }
    return redirect(url_for('admin_panel'))

# Cancelar edición de usuario
@app.route('/admin/clear_user_edit')
def clear_user_edit():
    session.pop('user_to_edit', None)
    return redirect(url_for('admin_panel'))

# Eliminar usuario
@app.route('/admin/delete_user/<int:user_id>', methods=['POST'])
def delete_user(user_id):
    if 'user_id' not in session or not session.get('is_admin'):
        flash('Acceso restringido.')
        return redirect(url_for('login'))
    
    user = User.query.get_or_404(user_id)
    
    # No permitir eliminar administradores
    if user.is_admin:
        flash('No se puede eliminar un usuario administrador.')
        return redirect(url_for('admin_panel'))
    
    # Eliminar datos relacionados del usuario
    UserAnswer.query.filter_by(user_id=user_id).delete()
    Achievement.query.filter_by(user_id=user_id).delete()
    # Notification.to_user is an Integer FK to users.id — use id, not username
    Notification.query.filter_by(to_user=user.id).delete()
    
    # Eliminar el usuario
    db.session.delete(user)
    db.session.commit()
    
    flash(f'Usuario "{user.username}" eliminado correctamente.')
    return redirect(url_for('admin_panel'))

# Bloquear/Desbloquear usuario
@app.route('/admin/toggle_block_user/<int:user_id>', methods=['POST'])
def toggle_block_user(user_id):
    if 'user_id' not in session or not session.get('is_admin'):
        flash('Acceso restringido.')
        return redirect(url_for('login'))
    
    user = User.query.get_or_404(user_id)
    
    # No permitir bloquear administradores
    if user.is_admin:
        flash('No se puede bloquear un usuario administrador.')
        return redirect(url_for('admin_panel'))
    
    # Cambiar estado de bloqueo
    user.is_blocked = not user.is_blocked
    db.session.commit()
    
    status = "bloqueado" if user.is_blocked else "desbloqueado"
    flash(f'Usuario "{user.username}" {status} correctamente.')
    return redirect(url_for('admin_panel'))

@app.route('/admin/assign_achievement', methods=['POST'])
def assign_achievement():
    if 'user_id' not in session or not session.get('is_admin'):
        flash('Acceso restringido.')
        return redirect(url_for('login'))
    user_id = request.form.get('user_id')
    badge_id = request.form.get('badge_id')
    if not user_id or not badge_id:
        flash('Usuario y logro requeridos.')
        return redirect(url_for('admin_panel'))
    # Evitar duplicados
    exists = Achievement.query.filter_by(user_id=user_id, badge_id=badge_id).first()
    if exists:
        flash('El usuario ya tiene este logro.')
        return redirect(url_for('admin_panel'))
    achievement = Achievement(user_id=user_id, badge_id=badge_id)
    db.session.add(achievement)
    db.session.commit()
    flash('Logro asignado correctamente.')
    return redirect(url_for('admin_panel'))

@app.route('/teacher/dashboard', methods=['GET', 'POST'])
def teacher_dashboard():
    if 'user_id' not in session:
        return redirect(url_for('login'))
    user = User.query.get(session['user_id'])
    from models import Game
    teacher = Teacher.query.filter_by(user_id=user.id).first()
    if not teacher:
        flash('No tienes perfil de profesor.')
        return redirect(url_for('dashboard'))
    quizzes = Quiz.query.filter_by(teacher_id=teacher.id).all()
    games = Game.query.filter_by(teacher_id=teacher.id).all()
    badges = Badge.query.all()
    achievements = Achievement.query.filter(Achievement.badge_id.in_([b.id for b in badges])).all()
    notifications = Notification.query.filter_by(to_user=user.id).order_by(Notification.date_sent.desc()).limit(10).all()

    # Alumnos inscritos al área del profesor
    students = User.query.filter_by(area=teacher.area, is_admin=False).all()

    # Puntajes de los alumnos en quizzes y juegos creados por el profesor
    quiz_ids = [q.id for q in quizzes]
    game_ids = [g.id for g in games]
    student_scores = []
    for student in students:
        quiz_points = 0
        if quiz_ids:
                quiz_points = db.session.query(
                    db.func.sum(
                        db.case((UserAnswer.is_correct == True, 10), else_=0)
                    )
                ).join(Question).filter(
                    UserAnswer.user_id==student.id,
                    Question.quiz_id.in_(quiz_ids)
                ).scalar() or 0
        student_scores.append({
            'student': student,
            'quiz_points': quiz_points
        })

    ia_questions = None
    ia_tema = None
    if request.method == 'POST' and 'tema' in request.form and 'cantidad' in request.form:
        tema = request.form['tema']
        cantidad = int(request.form['cantidad'])
        ia_tema = tema
        tipo_examen = request.form.get('tipo_examen', 'simple')  # 'simple' o 'opciones'
        ai_method = request.form.get('ai_method', 'offline')  # 'ollama', 'huggingface', 'offline'
        
        try:
            # Importar nuestro generador de IA local
            from ai_local import generate_local_exam
            
            ia_questions = generate_local_exam(tema, cantidad, tipo_examen, ai_method)
            
            # Si no se generaron preguntas, usar método offline como respaldo
            if not ia_questions:
                ia_questions = generate_local_exam(tema, cantidad, tipo_examen, 'offline')
                
        except Exception as e:
            # Fallback: generar preguntas básicas
            ia_questions = []
            for i in range(cantidad):
                if tipo_examen == 'opciones':
                    ia_questions.append({
                        'pregunta': f"Pregunta {i+1} sobre {tema}",
                        'opciones': ['Opción A', 'Opción B', 'Opción C', 'Opción D'],
                        'respuesta': 'A'
                    })
                else:
                    ia_questions.append(f"Pregunta {i+1} sobre {tema}")

    return render_template('teacher_dashboard.html',
                           user=user,
                           teacher=teacher,
                           quizzes=quizzes,
                           games=games,
                           badges=badges,
                           achievements=achievements,
                           notifications=notifications,
                           students=students,
                           student_scores=student_scores,
                           ia_questions=ia_questions,
                           ia_tema=ia_tema,
                           ia_tipo_examen=request.form.get('tipo_examen', 'simple'))

# === NUEVAS RUTAS PARA PROFESOR ===
@app.route('/teacher/create_quiz', methods=['POST'])
def teacher_create_quiz():
    if 'user_id' not in session:
        flash('Acceso restringido.')
        return redirect(url_for('login'))
    user = User.query.get(session['user_id'])
    teacher = Teacher.query.filter_by(user_id=user.id).first()
    if not teacher:
        flash('No tienes perfil de profesor.')
        return redirect(url_for('dashboard'))
    title = request.form['title'].strip()
    area = request.form['area'].strip()
    if not title or not area:
        flash('Todos los campos son obligatorios.')
        return redirect(url_for('teacher_dashboard'))
    new_quiz = Quiz(
        title=title,
        area=area,
        description=request.form.get('description', ''),
        created_by=user.id,
        teacher_id=teacher.id,
        created_at=datetime.utcnow()
    )
    db.session.add(new_quiz)
    db.session.commit()
    flash('Quiz creado correctamente.')
    return redirect(url_for('teacher_dashboard'))

@app.route('/teacher/create_game', methods=['POST'])
def teacher_create_game():
    if 'user_id' not in session:
        flash('Acceso restringido.')
        return redirect(url_for('login'))
    user = User.query.get(session['user_id'])
    teacher = Teacher.query.filter_by(user_id=user.id).first()
    if not teacher:
        flash('No tienes perfil de profesor.')
        return redirect(url_for('dashboard'))
    name = request.form.get('name', '').strip()
    type_ = request.form.get('type', '').strip()
    if not name or not type_:
        flash('Todos los campos son obligatorios.')
        return redirect(url_for('teacher_dashboard'))
    from models import Game
    new_game = Game(name=name, type=type_, teacher_id=teacher.id, created_at=datetime.utcnow())
    db.session.add(new_game)
    db.session.commit()
    flash('Juego creado correctamente.')
    return redirect(url_for('teacher_dashboard'))

@app.route('/teacher/add_question', methods=['POST'])
def teacher_add_question():
    if 'user_id' not in session:
        flash('Acceso restringido.')
        return redirect(url_for('login'))
    user = User.query.get(session['user_id'])
    teacher = Teacher.query.filter_by(user_id=user.id).first()
    if not teacher:
        flash('No tienes perfil de profesor.')
        return redirect(url_for('dashboard'))
    quiz_id = request.form['quiz_id']
    text = request.form['question_text'].strip()
    options = [
        request.form['option1'].strip(),
        request.form['option2'].strip(),
        request.form['option3'].strip(),
        request.form['option4'].strip()
    ]
    correct_index = int(request.form['correct_answer']) - 1
    if not text or any(not o for o in options) or correct_index not in range(4):
        flash('Todos los campos son obligatorios y la respuesta debe ser válida.')
        return redirect(url_for('teacher_dashboard'))
    correct_answer = options[correct_index]
    quiz = Quiz.query.get(quiz_id)
    if not quiz or quiz.teacher_id != teacher.id:
        flash('No tienes permisos para agregar preguntas a este quiz.')
        return redirect(url_for('teacher_dashboard'))
    new_question = Question(
        quiz_id=quiz_id,
        text=text,
        options=options,
        correct_answer=correct_answer
    )
    db.session.add(new_question)
    db.session.commit()
    flash('Pregunta agregada correctamente.')
    return redirect(url_for('teacher_dashboard'))
# Descargar examen en Word
@app.route('/download_exam_word', methods=['POST'])
def download_exam_word():
    from flask import send_file
    from io import BytesIO
    import json
    
    tema = request.form.get('tema', 'Examen')
    area = request.form.get('area', 'General')
    tipo_examen = request.form.get('tipo_examen', 'simple')
    preguntas_raw = request.form.get('preguntas', '')
    
    print(f"DEBUG - Tema: {tema}, Area: {area}, Tipo: {tipo_examen}")
    print(f"DEBUG - Preguntas raw length: {len(preguntas_raw)}")
    print(f"DEBUG - Preguntas raw start: {preguntas_raw[:100]}...")
    
    # Debug: mostrar todos los datos del form
    print("DEBUG - Todos los datos del formulario:")
    for key, value in request.form.items():
        if len(str(value)) > 100:
            print(f"  {key}: {str(value)[:100]}... (truncado)")
        else:
            print(f"  {key}: {value}")
    
    if not preguntas_raw:
        print("DEBUG - ERROR: No se recibieron preguntas en el formulario")
        return "Error: No se encontraron preguntas para descargar", 400
    
    try:
        # Limpiar el string JSON primero
        preguntas_clean = preguntas_raw.strip()
        
        # Verificar si es un JSON válido
        if preguntas_clean.startswith('[') and preguntas_clean.endswith(']'):
            # Intentar parseo directo
            preguntas = json.loads(preguntas_clean)
        elif preguntas_clean.startswith("'[") and preguntas_clean.endswith("]'"):
            # Remover comillas externas si las hay
            preguntas_clean = preguntas_clean[1:-1]
            preguntas = json.loads(preguntas_clean)
        else:
            # Intentar reparar comillas
            preguntas_fixed = preguntas_clean.replace("'", '"')
            preguntas = json.loads(preguntas_fixed)
        
        print(f"DEBUG - Preguntas parseadas exitosamente: {len(preguntas)} preguntas")
        if preguntas:
            print(f"DEBUG - Primera pregunta keys: {list(preguntas[0].keys()) if isinstance(preguntas[0], dict) else 'No es dict'}")
            
    except (json.JSONDecodeError, ValueError) as e:
        print(f"DEBUG - Error JSON específico: {e}")
        print(f"DEBUG - Posición del error: {getattr(e, 'pos', 'N/A')}")
        print(f"DEBUG - Raw content para debug: {repr(preguntas_raw[:100])}")
        
        # Intentar métodos alternativos de parseo más específicos
        try:
            # Método 1: Usar ast.literal_eval para estructuras de Python
            import ast
            preguntas = ast.literal_eval(preguntas_raw)
            print(f"DEBUG - Preguntas extraídas con ast: {len(preguntas)}")
        except (ValueError, SyntaxError) as e2:
            print(f"DEBUG - ast también falló: {e2}")
            
            try:
                # Método 2: Intentar reparar JSON manualmente
                repaired = preguntas_raw.replace("'", '"').replace('True', 'true').replace('False', 'false').replace('None', 'null')
                preguntas = json.loads(repaired)
                print(f"DEBUG - JSON reparado exitosamente: {len(preguntas)} preguntas")
            except json.JSONDecodeError as e3:
                print(f"DEBUG - Reparación JSON falló: {e3}")
                
                # Último recurso: generar preguntas desde el tema original
                print(f"DEBUG - Generando preguntas de respaldo desde tema: {tema}")
                try:
                    from ai_local import generate_local_exam
                    preguntas = generate_local_exam(tema, 5, tipo_examen, 'offline')
                    print(f"DEBUG - Preguntas de respaldo generadas: {len(preguntas)}")
                except Exception as e4:
                    print(f"DEBUG - Generación de respaldo falló: {e4}")
                    # Última opción: mensaje de error útil
                    preguntas = [{
                        'pregunta': f'ERROR: No se pudieron procesar las preguntas para {tema}. Intenta regenerar el examen.',
                        'opciones': [
                            'Ir al generador y crear un nuevo examen',
                            'Usar un tema diferente (ej: matematicas)',
                            'Verificar que el tema sea valido',
                            'Contactar al administrador'
                        ],
                        'respuesta': 'A'
                    }]
    
    print(f"DEBUG - Total preguntas finales a procesar: {len(preguntas)}")
    
    # Validar estructura de preguntas
    if preguntas and isinstance(preguntas[0], dict):
        print(f"DEBUG - Estructura correcta de pregunta detectada")
        if 'pregunta' in preguntas[0]:
            print(f"DEBUG - Pregunta válida: {preguntas[0]['pregunta'][:50]}...")
        else:
            print(f"DEBUG - Advertencia: falta clave 'pregunta' en estructura")
    else:
        print(f"DEBUG - Estructura de pregunta inesperada: {type(preguntas[0]) if preguntas else 'Lista vacía'}")
    
    try:
        try:
            from docx import Document
            print("DEBUG - Librería python-docx importada exitosamente")
        except ImportError as import_error:
            print(f"DEBUG - Error importando python-docx: {import_error}")
            print("DEBUG - Generando archivo de texto como alternativa")
            return download_exam_simple()
        
        doc = Document()
        doc.add_heading(f'Examen de {area}', 0)
        doc.add_paragraph(f'Tema: {tema}')
        doc.add_paragraph('Instrucciones: Responde las siguientes preguntas.')
        doc.add_paragraph('')  # Espacio
        
        if tipo_examen == 'opciones' and preguntas and isinstance(preguntas[0], dict):
            for i, q in enumerate(preguntas, 1):
                # Agregar pregunta
                p = doc.add_paragraph()
                p.add_run(f'{i}. {q.get("pregunta", "")}').bold = True
                
                # Agregar opciones
                opciones = q.get('opciones', [])
                for idx, opt in enumerate(opciones):
                    doc.add_paragraph(f"    {chr(97+idx)}) {opt}")
                
                # Agregar respuesta correcta
                respuesta = q.get('respuesta', '')
                if respuesta:
                    p_resp = doc.add_paragraph()
                    p_resp.add_run(f"Respuesta correcta: {respuesta}").italic = True
                
                doc.add_paragraph('')  # Espacio entre preguntas
        else:
            for i, q in enumerate(preguntas, 1):
                if isinstance(q, dict):
                    pregunta_text = q.get('pregunta', str(q))
                else:
                    pregunta_text = str(q)
                
                p = doc.add_paragraph()
                p.add_run(f'{i}. {pregunta_text}').bold = True
                doc.add_paragraph('')  # Espacio
        
        f = BytesIO()
        doc.save(f)
        f.seek(0)
        filename = f"Examen_{area}_{tema}.docx".replace(' ', '_').replace('/', '_')
        return send_file(f, as_attachment=True, download_name=filename, 
                        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
                        
    except ImportError:
        # Si no está disponible python-docx, crear un archivo de texto simple
        content = f"EXAMEN DE {area.upper()}\n"
        content += f"Tema: {tema}\n"
        content += "Instrucciones: Responde las siguientes preguntas.\n\n"
        
        if tipo_examen == 'opciones' and preguntas and isinstance(preguntas[0], dict):
            for i, q in enumerate(preguntas, 1):
                content += f"{i}. {q.get('pregunta', '')}\n"
                opciones = q.get('opciones', [])
                for idx, opt in enumerate(opciones):
                    content += f"    {chr(97+idx)}) {opt}\n"
                content += f"Respuesta correcta: {q.get('respuesta', '')}\n\n"
        else:
            for i, q in enumerate(preguntas, 1):
                if isinstance(q, dict):
                    pregunta_text = q.get('pregunta', str(q))
                else:
                    pregunta_text = str(q)
                content += f"{i}. {pregunta_text}\n\n"
        
        f = BytesIO(content.encode('utf-8'))
        f.seek(0)
        filename = f"Examen_{area}_{tema}.txt".replace(' ', '_').replace('/', '_')
        return send_file(f, as_attachment=True, download_name=filename, mimetype='text/plain')

# Descargar examen en PDF
@app.route('/download_exam_pdf', methods=['POST'])
def download_exam_pdf():
    from flask import send_file
    from io import BytesIO
    import json
    
    tema = request.form.get('tema', 'Examen')
    area = request.form.get('area', 'General')
    tipo_examen = request.form.get('tipo_examen', 'simple')
    preguntas_raw = request.form.get('preguntas', '')
    
    print(f"DEBUG PDF - Tema: {tema}, Area: {area}, Tipo: {tipo_examen}")
    print(f"DEBUG PDF - Preguntas raw length: {len(preguntas_raw)}")
    print(f"DEBUG PDF - Preguntas raw start: {preguntas_raw[:100]}...")
    
    # Debug: mostrar todos los datos del form para PDF
    print("DEBUG PDF - Todos los datos del formulario:")
    for key, value in request.form.items():
        if len(str(value)) > 100:
            print(f"  {key}: {str(value)[:100]}... (truncado)")
        else:
            print(f"  {key}: {value}")
    
    if not preguntas_raw:
        print("DEBUG PDF - ERROR: No se recibieron preguntas en el formulario")
        return "Error: No se encontraron preguntas para descargar", 400
    try:
        # Intentar decodificar JSON con diferentes métodos
        if preguntas_raw.startswith('[') and preguntas_raw.endswith(']'):
            # Parece ser JSON válido
            preguntas = json.loads(preguntas_raw)
        else:
            # Podría ser JSON mal formateado, intentar repararlo
            preguntas_raw_fixed = preguntas_raw.replace("'", '"')  # Cambiar comillas simples por dobles
            preguntas = json.loads(preguntas_raw_fixed)
        
        print(f"DEBUG PDF - Preguntas parseadas exitosamente: {len(preguntas)} preguntas")
        if preguntas:
            print(f"DEBUG PDF - Primera pregunta keys: {list(preguntas[0].keys()) if isinstance(preguntas[0], dict) else 'No es dict'}")
            
    except Exception as e:
        print(f"DEBUG PDF - Error parseando JSON: {e}")
        print(f"DEBUG PDF - Raw content preview: {preguntas_raw[:200]}...")
        
        # Intentar métodos alternativos de parseo
        try:
            # Método 1: Usar ast.literal_eval
            import ast
            preguntas = ast.literal_eval(preguntas_raw)
            print(f"DEBUG PDF - Preguntas extraídas con ast: {len(preguntas)}")
        except Exception as e2:
            print(f"DEBUG PDF - ast también falló: {e2}")
            
            # Método 2: Usar eval (más peligroso pero como último recurso)
            try:
                preguntas = eval(preguntas_raw)
                print(f"DEBUG PDF - Preguntas extraídas con eval: {len(preguntas)}")
            except Exception as e3:
                print(f"DEBUG PDF - eval también falló: {e3}")
                
                # Último recurso: mensaje de error descriptivo
                preguntas = [{
                    'pregunta': f'ERROR AL PROCESAR PREGUNTAS: JSON inválido. Error: {str(e)[:100]}',
                    'opciones': [
                        'Regenerar el examen con IA Local',
                        'Usar formato de preguntas simples', 
                        'Verificar generación de preguntas',
                        'Contactar administrador'
                    ],
                    'respuesta': 'A'
                }]
    
    print(f"DEBUG PDF - Total preguntas finales a procesar: {len(preguntas)}")
    
    # Validar estructura de preguntas
    if preguntas and isinstance(preguntas[0], dict):
        print(f"DEBUG PDF - Estructura correcta de pregunta detectada")
    else:
        print(f"DEBUG PDF - Estructura de pregunta inesperada: {type(preguntas[0]) if preguntas else 'Lista vacía'}")
    
    try:
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.pdfgen import canvas
            print("DEBUG - Librería reportlab importada exitosamente")
        except ImportError as import_error:
            print(f"DEBUG - Error importando reportlab: {import_error}")
            print("DEBUG - Generando archivo de texto como alternativa para PDF")
            return download_exam_simple()
        
        f = BytesIO()
        c = canvas.Canvas(f, pagesize=letter)
        width, height = letter
        
        # Título
        c.setFont('Helvetica-Bold', 16)
        c.drawString(40, height-50, f'Examen de {area}')
        
        # Información del examen
        c.setFont('Helvetica', 12)
        c.drawString(40, height-80, f'Tema: {tema}')
        c.drawString(40, height-100, 'Instrucciones: Responde las siguientes preguntas.')
        
        y = height-130
        
        if tipo_examen == 'opciones' and preguntas and isinstance(preguntas[0], dict):
            for i, q in enumerate(preguntas, 1):
                # Verificar si necesita nueva página
                if y < 100:
                    c.showPage()
                    y = height-50
                
                # Pregunta
                pregunta_text = q.get('pregunta', f'Pregunta {i}')
                c.setFont('Helvetica-Bold', 12)
                c.drawString(40, y, f'{i}. {pregunta_text}')
                y -= 25
                
                # Opciones
                c.setFont('Helvetica', 11)
                opciones = q.get('opciones', [])
                for idx, opt in enumerate(opciones):
                    if y < 50:
                        c.showPage()
                        y = height-50
                    c.drawString(60, y, f"{chr(97+idx)}) {opt}")
                    y -= 20
                
                # Respuesta correcta
                respuesta = q.get('respuesta', '')
                if respuesta:
                    if y < 50:
                        c.showPage()
                        y = height-50
                    c.setFont('Helvetica-Oblique', 10)
                    c.drawString(60, y, f"Respuesta correcta: {respuesta}")
                    y -= 25
                
                y -= 10  # Espacio extra entre preguntas
        else:
            for i, q in enumerate(preguntas, 1):
                if y < 80:
                    c.showPage()
                    y = height-50
                
                if isinstance(q, dict):
                    pregunta_text = q.get('pregunta', str(q))
                else:
                    pregunta_text = str(q)
                
                c.setFont('Helvetica-Bold', 12)
                c.drawString(40, y, f'{i}. {pregunta_text}')
                y -= 35
        
        c.save()
        f.seek(0)
        filename = f"Examen_{area}_{tema}.pdf".replace(' ', '_').replace('/', '_')
        return send_file(f, as_attachment=True, download_name=filename, mimetype='application/pdf')
        
    except ImportError:
        # Si no está disponible reportlab, crear un archivo de texto simple
        content = f"EXAMEN DE {area.upper()}\n"
        content += f"Tema: {tema}\n"
        content += "Instrucciones: Responde las siguientes preguntas.\n\n"
        
        if tipo_examen == 'opciones' and preguntas and isinstance(preguntas[0], dict):
            for i, q in enumerate(preguntas, 1):
                content += f"{i}. {q.get('pregunta', '')}\n"
                opciones = q.get('opciones', [])
                for idx, opt in enumerate(opciones):
                    content += f"    {chr(97+idx)}) {opt}\n"
                content += f"Respuesta correcta: {q.get('respuesta', '')}\n\n"
        else:
            for i, q in enumerate(preguntas, 1):
                if isinstance(q, dict):
                    pregunta_text = q.get('pregunta', str(q))
                else:
                    pregunta_text = str(q)
                content += f"{i}. {pregunta_text}\n\n"
        
        f = BytesIO(content.encode('utf-8'))
        f.seek(0)
        filename = f"Examen_{area}_{tema}.txt".replace(' ', '_').replace('/', '_')
        return send_file(f, as_attachment=True, download_name=filename, mimetype='text/plain')


# Ruta de descarga alternativa que siempre funciona
@app.route('/download_exam_simple', methods=['POST'])
def download_exam_simple():
    """Descarga simple que siempre funciona, en formato texto plano"""
    from flask import send_file
    from io import BytesIO
    import json
    from exam_utils import create_text_exam, debug_preguntas_info
    
    tema = request.form.get('tema', 'Examen')
    area = request.form.get('area', 'General')
    tipo_examen = request.form.get('tipo_examen', 'simple')
    preguntas_raw = request.form.get('preguntas', '')
    
    # Debug más detallado
    print(f"DEBUG SIMPLE - Tema: {tema}")
    print(f"DEBUG SIMPLE - Area: {area}")
    print(f"DEBUG SIMPLE - Tipo: {tipo_examen}")
    print(f"DEBUG SIMPLE - Preguntas raw length: {len(preguntas_raw)}")
    print(f"DEBUG SIMPLE - Preguntas raw preview: {preguntas_raw[:200]}...")
    
    debug_info = debug_preguntas_info(preguntas_raw)
    print(f"DEBUG SIMPLE - Info completa: {debug_info}")
    
    # Intentar parsear con manejo de errores específico
    preguntas = None
    try:
        if not preguntas_raw or preguntas_raw.strip() == '':
            raise ValueError("Preguntas raw está vacío")
        
        preguntas = json.loads(preguntas_raw)
        print(f"DEBUG SIMPLE - JSON parseado exitosamente: {len(preguntas)} preguntas")
        
        # Verificar que las preguntas son válidas y usar directamente si el tema es programación
        if isinstance(preguntas, list) and len(preguntas) > 0:
            primera_pregunta = preguntas[0]
            print(f"DEBUG SIMPLE - Primera pregunta: {primera_pregunta}")
            
            # Si el tema original es programación, usar las preguntas directamente
            tema_normalizado = tema.lower()
            if any(prog_word in tema_normalizado for prog_word in ['programacion', 'programación', 'programming', 'codigo', 'código']):
                print(f"DEBUG SIMPLE - Tema es programación, usando preguntas originales")
                # Las preguntas ya están en 'preguntas', no hacer nada más
            else:
                # Solo validar si NO es tema de programación explícito
                if isinstance(primera_pregunta, dict) and 'pregunta' in primera_pregunta:
                    pregunta_texto = primera_pregunta['pregunta'].lower()
                    es_programacion = any(palabra in pregunta_texto for palabra in 
                                        ['función', 'variable', 'código', 'programa', 'algoritmo', 'python', 
                                         'javascript', 'java', 'php', 'c++', 'c#', 'operador', 'sintaxis', 'tipo de dato',
                                         'patrón', 'patron', 'diseño', 'diseno', 'herencia', 'polimorfismo',
                                         'encapsulación', 'encapsulacion', 'recursión', 'recursion', 'iteración', 
                                         'iteracion', 'complejidad', 'framework', 'biblioteca', 'api', 'debugging',
                                         'clase', 'objeto', 'método', 'metodo', 'estructura', 'datos', 'arrays',
                                         'listas', 'pilas', 'colas', 'árboles', 'arboles', 'grafos', 'hash',
                                         'implementa', 'utiliza', 'diferencia', 'significa', 'ventajas'])
                    print(f"DEBUG SIMPLE - Es pregunta de programación: {es_programacion}")
                    
                    if not es_programacion:
                        print("DEBUG SIMPLE - ADVERTENCIA: Las preguntas no parecen ser de programación")
                        preguntas = None  # Forzar fallback
            
        else:
            raise ValueError(f"Formato de preguntas inválido: {type(preguntas)}")
            
    except json.JSONDecodeError as e:
        print(f"DEBUG SIMPLE - Error JSON Decode: {e}")
        print(f"DEBUG SIMPLE - JSON problemático: {preguntas_raw}")
        preguntas = None
    except Exception as e:
        print(f"DEBUG SIMPLE - Error general: {e}")
        preguntas = None
    
    # Si hay problemas, generar preguntas de programación reales como fallback
    if not preguntas:
        print("DEBUG SIMPLE - Usando fallback - generando preguntas de programación")
        from ai_local import generate_local_exam
        
        try:
            # Intentar generar preguntas reales de programación como fallback
            tema_normalizado = tema.lower()
            if 'programacion' in tema_normalizado or 'programming' in tema_normalizado:
                preguntas = generate_local_exam("programacion", 5, tipo_examen, "offline")
                print(f"DEBUG SIMPLE - Fallback con IA: {len(preguntas)} preguntas generadas")
            else:
                # Solo usar preguntas genéricas si realmente no es programación
                if tipo_examen == 'opciones':
                    preguntas = []
                    for i in range(5):  # 5 preguntas de ejemplo
                        preguntas.append({
                            'pregunta': f'Pregunta {i+1} sobre {tema}: ¿Cuál es un concepto importante en esta materia?',
                            'opciones': [
                                f'Concepto A relacionado con {tema}',
                                f'Concepto B relacionado con {tema}',
                                f'Concepto C relacionado con {tema}',
                                f'Concepto D relacionado con {tema}'
                            ],
                            'respuesta': chr(65 + (i % 4))  # A, B, C, D rotativamente
                        })
                else:
                    preguntas = [
                        f'¿Cuáles son los fundamentos básicos de {tema}?',
                        f'¿Cómo se aplica {tema} en la práctica?',
                        f'¿Qué ventajas ofrece estudiar {tema}?',
                        f'¿Cuáles son los principales desafíos en {tema}?',
                        f'¿Cómo ha evolucionado {tema} en los últimos años?'
                    ]
                print(f"DEBUG SIMPLE - Fallback genérico: {len(preguntas)} preguntas")
        except Exception as fallback_error:
            print(f"DEBUG SIMPLE - Error en fallback: {fallback_error}")
            # Último recurso: preguntas muy básicas
            preguntas = [f"Pregunta básica {i} sobre {tema}" for i in range(1, 6)]
    
    # Crear contenido del examen
    content = create_text_exam(tema, area, tipo_examen, preguntas)
    
    # Crear archivo
    f = BytesIO(content.encode('utf-8'))
    f.seek(0)
    filename = f"Examen_{area}_{tema}_simple.txt".replace(' ', '_').replace('/', '_')
    
    print(f"DEBUG SIMPLE - Archivo creado: {filename}, tamaño: {len(content)} chars")
    
    return send_file(f, as_attachment=True, download_name=filename, mimetype='text/plain')


@app.route('/teacher/editar_quiz/<int:quiz_id>', methods=['GET', 'POST'])
def editar_quiz(quiz_id):
    if 'user_id' not in session:
        return redirect(url_for('login'))
    user = User.query.get(session['user_id'])
    quiz = Quiz.query.get_or_404(quiz_id)
    # Solo permitir editar si el usuario es el profesor creador
    teacher = Teacher.query.filter_by(user_id=user.id).first()
    if not teacher or quiz.teacher_id != teacher.id:
        flash('No tienes permisos para editar este quiz.')
        return redirect(url_for('teacher_dashboard'))
    if request.method == 'POST':
        quiz.title = request.form['title'].strip()
        quiz.area = request.form['area'].strip()
        quiz.description = request.form.get('description', '').strip()
        db.session.commit()
        flash('Quiz actualizado correctamente.')
        return redirect(url_for('teacher_dashboard'))
    return render_template('edit_quiz.html', quiz=quiz)

@app.route('/teacher/eliminar_quiz/<int:quiz_id>', methods=['POST'])
def eliminar_quiz(quiz_id):
    if 'user_id' not in session:
        return redirect(url_for('login'))
    user = User.query.get(session['user_id'])
    quiz = Quiz.query.get_or_404(quiz_id)
    teacher = Teacher.query.filter_by(user_id=user.id).first()
    if not teacher or quiz.teacher_id != teacher.id:
        flash('No tienes permisos para eliminar este quiz.')
        return redirect(url_for('teacher_dashboard'))
    db.session.delete(quiz)
    db.session.commit()
    flash('Quiz eliminado correctamente.')
    return redirect(url_for('teacher_dashboard'))

@app.route('/teacher/editar_juego/<int:game_id>', methods=['GET', 'POST'])
def editar_juego(game_id):
    if 'user_id' not in session:
        return redirect(url_for('login'))
    user = User.query.get(session['user_id'])
    from models import Game
    game = Game.query.get_or_404(game_id)
    teacher = Teacher.query.filter_by(user_id=user.id).first()
    if not teacher or game.teacher_id != teacher.id:
        flash('No tienes permisos para editar este juego.')
        return redirect(url_for('teacher_dashboard'))
    if request.method == 'POST':
        game.name = request.form['name'].strip()
        game.description = request.form['description'].strip()
        game.type = request.form['type'].strip()
        game.rules = request.form['rules'].strip()
        db.session.commit()
        flash('Juego actualizado correctamente.')
        return redirect(url_for('teacher_dashboard'))
    return render_template('edit_game.html', game=game)

@app.route('/teacher/eliminar_juego/<int:game_id>', methods=['POST'])
def eliminar_juego(game_id):
    if 'user_id' not in session:
        return redirect(url_for('login'))
    user = User.query.get(session['user_id'])
    from models import Game
    game = Game.query.get_or_404(game_id)
    teacher = Teacher.query.filter_by(user_id=user.id).first()
    if not teacher or game.teacher_id != teacher.id:
        flash('No tienes permisos para eliminar este juego.')
        return redirect(url_for('teacher_dashboard'))
    db.session.delete(game)
    db.session.commit()
    flash('Juego eliminado correctamente.')
    return redirect(url_for('teacher_dashboard'))

# Editar pregunta (profesor)
@app.route('/teacher/editar_pregunta/<int:question_id>', methods=['GET', 'POST'])
def editar_pregunta(question_id):
    if 'user_id' not in session:
        return redirect(url_for('login'))
    user = User.query.get(session['user_id'])
    question = Question.query.get_or_404(question_id)
    quiz = Quiz.query.get(question.quiz_id)
    teacher = Teacher.query.filter_by(user_id=user.id).first()
    if not teacher or quiz.teacher_id != teacher.id:
        flash('No tienes permisos para editar esta pregunta.')
        return redirect(url_for('teacher_dashboard'))
    if request.method == 'POST':
        question.text = request.form['question_text'].strip()
        question.options = [
            request.form['option1'].strip(),
            request.form['option2'].strip(),
            request.form['option3'].strip(),
            request.form['option4'].strip()
        ]
        correct_index = int(request.form['correct_answer']) - 1
        if correct_index not in range(4):
            flash('La respuesta correcta debe ser válida.')
            return redirect(url_for('editar_pregunta', question_id=question.id))
        question.correct_answer = question.options[correct_index]
        db.session.commit()
        flash('Pregunta actualizada correctamente.')
        return redirect(url_for('teacher_dashboard'))
    return render_template('edit_question.html', question=question)

# Eliminar pregunta (profesor)
@app.route('/teacher/eliminar_pregunta/<int:question_id>', methods=['POST'])
def eliminar_pregunta(question_id):
    if 'user_id' not in session:
        return redirect(url_for('login'))
    user = User.query.get(session['user_id'])
    question = Question.query.get_or_404(question_id)
    quiz = Quiz.query.get(question.quiz_id)
    teacher = Teacher.query.filter_by(user_id=user.id).first()
    if not teacher or quiz.teacher_id != teacher.id:
        flash('No tienes permisos para eliminar esta pregunta.')
        return redirect(url_for('teacher_dashboard'))
    db.session.delete(question)
    db.session.commit()
    flash('Pregunta eliminada correctamente.')
    return redirect(url_for('teacher_dashboard'))

# Editar pregunta (profesor)
@app.route('/teacher/edit_question/<int:question_id>', methods=['GET', 'POST'])
def edit_question(question_id):
    if 'user_id' not in session:
        return redirect(url_for('login'))
    user = User.query.get(session['user_id'])
    teacher = Teacher.query.filter_by(user_id=user.id).first()
    question = Question.query.get_or_404(question_id)
    quiz = Quiz.query.get(question.quiz_id)
    if not teacher or quiz.teacher_id != teacher.id:
        flash('No tienes permisos para editar esta pregunta.')
        return redirect(url_for('teacher_dashboard'))
    if request.method == 'POST':
        text = request.form['question_text'].strip()
        options = [
            request.form['option1'].strip(),
            request.form['option2'].strip(),
            request.form['option3'].strip(),
            request.form['option4'].strip()
        ]
        correct_index = int(request.form['correct_answer']) - 1
        if not text or any(not o for o in options) or correct_index not in range(4):
            flash('Todos los campos son obligatorios y la respuesta debe ser válida.')
            return render_template('edit_question.html', question=question)
        question.text = text
        question.options = options
        question.correct_answer = options[correct_index]
        db.session.commit()
        flash('Pregunta actualizada correctamente.')
        return redirect(url_for('teacher_dashboard'))
    return render_template('edit_question.html', question=question)

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    socketio.run(app, host='0.0.0.0', port=port, debug=False)

